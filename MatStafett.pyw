"""
Project: Food Relay
Author: Patrik Viljehav
Version: 1.0
URL: https://github.com/pvilje/matstafett
Description:    Select an Excel (.xlsx) document.
                The script will calculate a setup where all participants
                will:
                    * Get a certain part of the dinner to prepare.
                    * Meet new people at every part of the dinner
                    * Save the result in the same file type as was submitted to the script (this can be changed)

Non standard Dependencies:
                * openpyxl verified to work on versions: 2.4.0 - 2.4.6
                * docx verified to work on version: 0.8.6

"""

# Import needed modules
import os
import random
import csv
import tkinter
from tkinter import filedialog
from tkinter import messagebox
from collections import deque
import openpyxl
from openpyxl.styles import *
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Constants
# =========
OPENPYXL_VERSION = "2.4"
PYTHON_DOCX_VERSION = "0.8"


def rotate(l, n):
    """
    Rotates list "l" "n" times.
    :param l: the list to rotate
    :param n: times to rotate
    :return: a rotated list object
    """
    temp = deque(l)
    temp.rotate(n)
    return list(temp)


class Hmi:
    def __init__(self, parent, language="eng", csv_file="lang.csv"):
        """
        Initialize the class
        :param parent: The master tk window
        :param language: the GUI language, defaults to english.
        :param csv_file: The csv file containing the language strings
        """
        # Set up variables
        # ================
        self.gui_language = ""
        self.csv_file = ""
        self.file_type = ""
        self.file_name = None
        self.file_path = None
        self.list_supported_file_types = [("Excel", "*.xlsx")]
        self.list_participants = []
        self.list_sorted_participants = []
        self.groups_starter = []
        self.groups_main = []
        self.groups_desert = []
        self.list_rand_index = []
        self.num_groups = None
        self.lang = {}
        self.host_s = []
        self.host_m = []
        self.host_d = []
        self.guest_s_1 = []
        self.guest_s_2 = []
        self.guest_m_1 = []
        self.guest_m_2 = []
        self.guest_d_1 = []
        self.guest_d_2 = []
        self.sorted_result = []

        # Get language pack
        # =================
        # Make sure it is a csv file. Allow different formats, but show a warning that it seems erroneous
        # This can not be translated since the language file obviously does not work.
        if not csv_file.lower().endswith(".csv"):
            if not messagebox.askyesno("Invalid language file",
                                       "Invalid File format for language file, expected '.csv', got .{}\n\n"
                                       "Do you want to try to use the supplied file anyway?\n"
                                       "({})".format(
                                           csv_file.lower().split(".")[-1], csv_file)):
                quit()
        # All is fine, get the phrases for the selected language
        self.csv_file = csv_file
        self.gui_language = language
        self.get_lang()
        parent.title(self.lang["label_title"])

        # ================
        # tkinter stuff...
        # ================

        # Main frames and labels
        # ======================
        self.l_title = tkinter.Label(parent, text=self.lang["label_title"])
        self.f_input = tkinter.Frame(parent, pady=10, padx=10)
        self.f_output = tkinter.Frame(parent, pady=10, padx=10)
        self.l_options = tkinter.Label(self.f_input, text=self.lang["label_options"])

        # TK Variables
        # ============
        self.sv_filename = tkinter.StringVar()
        self.iv_new_year_same_lineup = tkinter.IntVar()
        self.iv_generate_letters = tkinter.IntVar()

        # Input widgets
        # =============
        self.e_filename = tkinter.Entry(self.f_input,
                                        textvariable=self.sv_filename,
                                        width=60)
        self.b_select_file = tkinter.Button(self.f_input,
                                            text=self.lang["file_select"],
                                            command=self.select_file)
        self.b_run = tkinter.Button(self.f_input,
                                    text=self.lang["button_run"],
                                    command=self.generate_result,
                                    state=tkinter.DISABLED,
                                    height=3, width=10)
        self.cb_new_year_same_lineup = tkinter.Checkbutton(self.f_input,
                                                           text=self.lang["button_same_lineup"],
                                                           variable=self.iv_new_year_same_lineup)
        self.cb_generate_letters = tkinter.Checkbutton(self.f_input,
                                                       text=self.lang["button_generate_letters"],
                                                       variable=self.iv_generate_letters)

        # Output widgets
        # ==============
        self.t_output = tkinter.Text(self.f_output,
                                     height=10, width=70,
                                     state=tkinter.DISABLED)

        # Add colors and scroll-bars to the output widget
        # ===============================================
        log_colors = ["black", "green", "blue", "red"]
        for color in log_colors:
            self.t_output.tag_config(color, foreground=color)
        self.scroll_x_output = tkinter.Scrollbar(self.f_output, command=self.t_output.xview, orient=tkinter.HORIZONTAL)
        self.scroll_y_output = tkinter.Scrollbar(self.f_output, command=self.t_output.yview)
        self.t_output.configure(yscrollcommand=self.scroll_y_output.set, xscrollcommand=self.scroll_x_output.set)

        # Check openpyxl and docx version
        # ======================
        if not openpyxl.__version__.startswith(OPENPYXL_VERSION):
            tkinter.messagebox.showwarning(self.lang["warning_unexpected_version_title"],
                                           self.lang["warning_openpyxl_version"]
                                           .format(OPENPYXL_VERSION, openpyxl.__version__))
        if not docx.__version__.startswith(PYTHON_DOCX_VERSION):
            tkinter.messagebox.showwarning(self.lang["warning_unexpected_version_title"],
                                           self.lang["warning_docx_version"]
                                           .format(PYTHON_DOCX_VERSION, docx.__version__))

    def draw_main(self):
        """
        Draw the main window
        """
        # Main title
        self.l_title.grid(row=0, column=0)

        # input frame
        self.f_input.grid(row=1, column=0)

        self.e_filename.grid(row=0, column=0, columnspan=2)
        self.b_select_file.grid(row=0, column=2, sticky=tkinter.W, padx=15)
        self.l_options.grid(row=1, column=0, sticky=tkinter.W)
        self.cb_new_year_same_lineup.grid(row=2, column=0, sticky=tkinter.W)
        self.cb_generate_letters.grid(row=3, column=0, sticky=tkinter.W)
        self.b_run.grid(row=4, column=2)

        # output frame
        self.f_output.grid(row=2, column=0)
        self.t_output.grid(row=0, column=0, columnspan=2)
        self.scroll_x_output.grid(row=1, column=0, columnspan=2, sticky=tkinter.E + tkinter.W)
        self.scroll_y_output.grid(row=0, column=2, sticky=tkinter.N + tkinter.S)

        # Preselect
        self.iv_generate_letters.set(1)

    def select_file(self):
        """
        Method to open a file dialog and validate the file type.
        """
        # Select file dialog.
        file = filedialog.askopenfilename(title=self.lang["dialog_select_file"],
                                          initialdir=os.curdir,
                                          filetypes=self.list_supported_file_types)
        # populate String-var and entry.
        self.sv_filename.set(file)
        self.e_filename = self.sv_filename.get()

        # Validate file
        # =============
        file_type_ok = False
        if len(file) > 0:
            if file.endswith(".xlsx"):
                self.file_type = ".xlsx"
                file_type_ok = True
            else:
                self.log_output(self.lang["error_file_types"], "red")
        else:
            self.log_output(self.lang["error_no_file_selected"], "red")

        # Activate / deactivate run-button depending on file validation result
        # ====================================================================
        if file_type_ok:
            self.b_run.configure(state=tkinter.ACTIVE)
            self.file_path, self.file_name = os.path.split(file)
            self.log_output("{}: {}".format(self.lang["file_selected"], file))

        else:
            self.b_run.configure(state=tkinter.DISABLED)
            self.file_name = None
            self.file_path = None
            self.file_type = ""

    def read_file_contents(self):
        """
        Read the participant list in the selected file.
        :return result: 0 - OK, 1 - Unsupported File
        """
        # prepare to get all participants
        self.list_participants = []
        file = os.path.join(self.file_path, self.file_name)
        self.log_output(self.lang["file_reading_{}".format(self.file_type)])

        # If it is a xlsx file check column A for participants.
        # =====================================================
        if self.file_type == ".xlsx":
            wb = openpyxl.load_workbook(file)
            ws = wb.get_sheet_by_name(wb.sheetnames[0])
            max_rows = ws.max_row
            for number, row in enumerate(ws["A1:A{}".format(max_rows)], start=1):
                for cell in row:
                    if cell.value is not None:
                        # address column
                        address = ws["B{}".format(number)].value
                        # allergies column
                        allergies = ws["C{}".format(number)].value
                        # [name, address, allergies]
                        self.list_participants.append([cell.value, address, allergies])
            return 0
        else:
            # Should not be possible to end up here, but just in case...
            return 1

    def create_lineup(self):
        """
        Create a lineup for the food rally.
        """
        # Reset variables
        # ===============
        base_index = 0
        offset_1 = 1
        offset_2 = 2
        self.host_s = []
        self.host_m = []
        self.host_d = []
        self.guest_s_1 = []
        self.guest_s_2 = []
        self.guest_m_1 = []
        self.guest_m_2 = []
        self.guest_d_1 = []
        self.guest_d_2 = []
        self.sorted_result = []

        # loop through all groups, using three indexes for the group lists
        # =================================================================
        while base_index < self.num_groups:
            if offset_1 >= self.num_groups:
                offset_1 = 0
            if offset_2 >= self.num_groups:
                offset_2 = 0
            # starters
            self.host_s.append(self.groups_starter[base_index])
            self.guest_s_1.append(self.groups_main[base_index])
            self.guest_s_2.append(self.groups_desert[base_index])

            # main course
            self.host_m.append(self.groups_main[offset_1])
            self.guest_m_1.append(self.groups_starter[base_index])
            self.guest_m_2.append(self.groups_desert[offset_2])

            # desert
            self.host_d.append(self.groups_desert[offset_1])
            self.guest_d_1.append(self.groups_starter[base_index])
            self.guest_d_2.append(self.groups_main[offset_2])

            base_index += 1
            offset_1 += 1
            offset_2 += 1

    def save_to_file(self):
        """
        Save the generated list to a file, Grouped and neatly ordered
        """
        self.log_output(self.lang["progress_gen_route"])

        # Save an .xlsx file
        # =================
        if self.file_type == ".xlsx":
            # Save to a new file, don't mess with the source.
            # check if this is the first generated result
            if os.path.isfile(os.path.join(self.file_path, "{}_{}".format(self.lang["file_name_result"],
                                                                          self.file_name))):
                # File already existed, generate a new filename (add a number until a not used name is found)
                file_no = 2
                while os.path.isfile(
                        os.path.join(
                            self.file_path, "{}_{}_{}".format(self.lang["file_name_result"], str(file_no),
                                                              self.file_name))):
                    file_no += 1
                file = os.path.join(
                    self.file_path, "{}_{}_{}".format(self.lang["file_name_result"], str(file_no), self.file_name))
            else:
                # This is the first generated result.
                file = os.path.join(self.file_path, "{}_{}".format(self.lang["file_name_result"], self.file_name))

            # Open Excel workbook
            # ===================
            wb = openpyxl.Workbook()

            # Sheet 1 list of participants
            # ============================

            ws = wb.active
            ws.title = self.lang["excel_sheet_name_one"]

            for line_no, participant in enumerate(self.list_sorted_participants):
                ws["A{}".format(line_no + 1)] = participant[0]
                if participant[1] is not None:
                    ws["B{}".format(line_no + 1)] = participant[1]
                if participant[2] is not None:
                    ws["C{}".format(line_no + 1)] = participant[2]

            ws2 = wb.create_sheet(self.lang["excel_sheet_name_two"])

            # Sheet 2 Generated result
            # ============================

            # Setup styles
            # ============
            ws2.column_dimensions["A"].width = 34
            ws2.column_dimensions["C"].width = 34
            ws2.column_dimensions["D"].width = 34
            ws2.column_dimensions["E"].width = 34
            h1 = NamedStyle(name="h1", font=Font(size=15, bold=True, color="1f497d"),
                            border=Border(bottom=Side(style="thick", color="4f81bd")))
            h1_center = NamedStyle(name="h1_center",
                                   font=Font(size=15, bold=True, color="1f497d"),
                                   border=Border(bottom=Side(style="thick", color="4f81bd")),
                                   alignment=Alignment(horizontal="center"))
            h2 = NamedStyle(name="h2", font=Font(size=13, bold=True, color="1f497d"),
                            border=Border(bottom=Side(style="thick", color="a7bfde")))
            h2_center = NamedStyle(name="h2_center",
                                   font=Font(size=13, bold=True, color="1f497d"),
                                   border=Border(bottom=Side(style="thick", color="a7bfde")),
                                   alignment=Alignment(horizontal="center"))

            # Column A, summary of participants
            # =================================

            # Summary header
            ws2["A1"] = "{}".format(self.lang["excel_summary"])
            ws2["A1"].style = h1

            # Starter hosts
            ws2["A2"] = "{} {}:".format(self.lang["excel_host"], self.lang["excel_starter"])
            ws2["A2"].style = h2
            row = 3
            for name in self.host_s:
                ws2["A{}".format(row)] = name[0]
                row += 1
            row += 2

            # Main course hosts
            ws2["A{}".format(row)] = "{} {}:".format(self.lang["excel_host"], self.lang["excel_main_course"])
            ws2["A{}".format(row)].style = h2
            row += 1
            for name in self.host_m:
                ws2["A{}".format(row)] = name[0]
                row += 1
            row += 2

            # Desert hosts
            ws2["A{}".format(row)] = "{} {}:".format(self.lang["excel_host"], self.lang["excel_desert"])
            ws2["A{}".format(row)].style = h2
            row += 1
            for name in self.host_d:
                ws2["A{}".format(row)] = name[0]
                row += 1

            # Column C, D, E: results
            # =================================
            # Starters
            ws2.merge_cells("C1:E1")
            ws2["C1"] = self.lang["excel_starter"]
            ws2["C1"].style = h1_center
            ws2["D1"].style = h1_center  # needed for the border, even though the cells are merged
            ws2["E1"].style = h1_center  # needed for the border, even though the cells are merged
            ws2["C2"] = self.lang["excel_host"]
            ws2["D2"] = self.lang["excel_guest"]
            ws2["E2"] = self.lang["excel_guest"]
            ws2["C2"].style = h2_center
            ws2["D2"].style = h2_center
            ws2["E2"].style = h2_center
            row = 3
            for index, host in enumerate(self.host_s):
                ws2["C{}".format(row)] = host[0]
                ws2["D{}".format(row)] = self.guest_s_1[index][0]
                ws2["E{}".format(row)] = self.guest_s_2[index][0]
                row += 1

            # Main Course
            row += 1
            ws2.merge_cells("C{}:E{}".format(row, row))
            ws2["C{}".format(row)] = self.lang["excel_main_course"]
            ws2["C{}".format(row)].style = h1_center
            ws2["D{}".format(row)].style = h1_center  # needed for the border, even though the cells are merged
            ws2["E{}".format(row)].style = h1_center  # needed for the border, even though the cells are merged
            row += 1
            ws2["C{}".format(row)] = self.lang["excel_host"]
            ws2["D{}".format(row)] = self.lang["excel_guest"]
            ws2["E{}".format(row)] = self.lang["excel_guest"]
            ws2["C{}".format(row)].style = h2_center
            ws2["D{}".format(row)].style = h2_center
            ws2["E{}".format(row)].style = h2_center
            row += 1
            for index, host in enumerate(self.host_m):
                ws2["C{}".format(row)] = host[0]
                ws2["D{}".format(row)] = self.guest_m_1[index][0]
                ws2["E{}".format(row)] = self.guest_m_2[index][0]
                row += 1

            # Desert
            row += 1
            ws2.merge_cells("C{}:E{}".format(row, row))
            ws2["C{}".format(row)] = self.lang["excel_desert"]
            ws2["C{}".format(row)].style = h1_center
            ws2["D{}".format(row)].style = h1_center  # needed for the border, even though the cells are merged
            ws2["E{}".format(row)].style = h1_center  # needed for the border, even though the cells are merged
            row += 1
            ws2["C{}".format(row)] = self.lang["excel_host"]
            ws2["D{}".format(row)] = self.lang["excel_guest"]
            ws2["E{}".format(row)] = self.lang["excel_guest"]
            ws2["C{}".format(row)].style = h2_center
            ws2["D{}".format(row)].style = h2_center
            ws2["E{}".format(row)].style = h2_center
            row += 1
            for index, host in enumerate(self.host_d):
                ws2["C{}".format(row)] = host[0]
                ws2["D{}".format(row)] = self.guest_d_1[index][0]
                ws2["E{}".format(row)] = self.guest_d_2[index][0]
                row += 1

            # Save! (also closes the file)
            try:
                wb.save(filename=file)
            except PermissionError:
                self.log_output(self.lang["error_save"], "red")
            except FileNotFoundError:
                self.log_output(self.lang["error_save"], "red")
            else:
                self.log_output(self.lang["progress_done"])
                self.log_output("{} \n{}".format(self.lang["progress_saved_to"], file))

        else:
            # Should not be possible to end up here, but just in case...
            self.log_output(self.lang["error_file_types"], "red")

    def validate_number_of_participants(self):
        """
        Makes sure the number of participants is:
            * a factor of 3.
            * at least 9
        Raises ValueError if not.
        """
        # Get all the participants:
        # =========================
        list_participants = self.list_participants

        # Verify that it is an ok number of participants
        # ==============================================
        if len(list_participants) < 9:
            # Must be 9 or more
            raise ValueError(self.lang["error_less_than_nine"])
        elif len(list_participants) % 3 != 0:
            # Must be a factor of 3
            raise ValueError(self.lang["error_number_participants"])
        else:
            # All is ok, go on.
            self.num_groups = int(len(list_participants) / 3)
            self.log_output("{}: {}".format(self.lang["progress_found_participants"], len(list_participants)))

    def generate_random_index(self):
        """
        Generate a list containing integers from 1 - amount of participants.
        Then shuffle it around so they don't all come in order...
        """
        self.log_output(self.lang["progress_gen_rand_list"])
        i = 0
        self.list_rand_index = []
        while i < len(self.list_participants):
            self.list_rand_index.append(i)
            i += 1
        random.shuffle(self.list_rand_index)

    def sort_participants(self):
        """
        Generate a sorted... sort of unsorted... list of members based on random numbers.
        """
        self.log_output(self.lang["progress_sort_unsort"])
        self.list_sorted_participants = []

        # Is this based on a previous result or not?
        if self.iv_new_year_same_lineup.get():
            # Get the list of participants, but shift all the hosts so that no one hosts the same part as last time.
            self.groups_desert = self.list_participants[0:self.num_groups]
            self.groups_starter = self.list_participants[self.num_groups:self.num_groups * 2]
            self.groups_main = self.list_participants[self.num_groups * 2:self.num_groups * 3]

            # Rotate the lists to make sure no one meets the same participants as last round.
            self.groups_starter = rotate(self.groups_starter, 0)
            self.groups_main = rotate(self.groups_main, 4)
            self.groups_desert = rotate(self.groups_desert, 8)

            # copy all participants to a sorted list.
            self.list_sorted_participants += self.groups_starter
            self.list_sorted_participants += self.groups_main
            self.list_sorted_participants += self.groups_desert
        else:
            for index in self.list_rand_index:
                self.list_sorted_participants.append(self.list_participants[index])

        # create three equal sized lists containing all participants
        # ==========================================================
            self.groups_starter = self.list_sorted_participants[0:self.num_groups]
            self.groups_main = self.list_sorted_participants[self.num_groups:self.num_groups * 2]
            self.groups_desert = self.list_sorted_participants[self.num_groups * 2:self.num_groups * 3]

    def generate_docx_letters(self):
        """
        Generates a docx file containing all the letters.
        """
        # Base filename.
        file_name = self.file_name.split(".")[0] + self.lang["word_file_name_letter"]
        file_ext = ".docx"
        # Find a free file name
        # =====================
        if os.path.isfile(os.path.join(self.file_path, "{}{}".format(file_name, file_ext))):
            # File already existed, generate a new filename (add a number until a not used name is found)
            file_no = 2
            while os.path.isfile(
                    os.path.join(
                        self.file_path, "{}{}{}".format(file_name, str(file_no), file_ext))):
                file_no += 1
            file = os.path.join(
                self.file_path, "{}{}{}".format(file_name, str(file_no), file_ext))
        else:
            file = os.path.join(self.file_path, "{}{}".format(file_name, file_ext))

        # Create a docx object
        document = Document()

        #  Create a page in the document to tell each participant what to host and where to have starters.
        # Starters
        document = self.print_to_word(document, "starter", self.host_s)
        # Main course
        document = self.print_to_word(document, "main_course", self.host_m)
        # Desert
        document = self.print_to_word(document, "desert", self.host_d)

        # Add pages for where the participants should go next:
        document = self.print_to_word(document, "where_to_go", None)

        # Style the document!
        # ===================
        # Body text 3
        document.styles["Body Text 3"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.styles["Body Text 3"].font.name = "Lucida Calligraphy"
        document.styles["Body Text 3"].font.size = Pt(12)
        document.styles["Body Text 3"].font.underline = True
        document.styles["Body Text 3"].hidden = False
        document.styles["Body Text 3"].quick_style = True
        # Body text 2
        document.styles["Body Text 2"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.styles["Body Text 2"].font.name = "Calibri"
        document.styles["Body Text 2"].font.size = Pt(12)
        document.styles["Body Text 2"].hidden = False
        document.styles["Body Text 2"].quick_style = True
        document.styles["Body Text 2"].paragraph_format.space_before = Pt(12)
        document.styles["Body Text 2"].paragraph_format.space_after = Pt(12)
        document.styles["Body Text 2"].paragraph_format.line_spacing = Pt(12)
        # Body Text
        document.styles["Body Text"].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.styles["Body Text"].font.name = "Calibri"
        document.styles["Body Text"].font.size = Pt(12)
        document.styles["Body Text"].font.italic = True
        document.styles["Body Text"].hidden = False
        document.styles["Body Text"].quick_style = True
        document.styles["Body Text"].paragraph_format.space_before = Pt(12)
        document.styles["Body Text"].paragraph_format.space_after = Pt(12)
        document.styles["Body Text"].paragraph_format.line_spacing = Pt(12)

        # Save
        # ====
        document.save(file)
        self.log_output("{} \n{}".format(self.lang["progress_saved_to"], file))

    def print_to_word(self, doc, part, content):
        """
        Prints specified content to the word document
        :param doc: docx document object
        :param part: part of the meal
        :param content: the contents
        :return: the modiefied document
        """
        def find_index(needle, haystack):
            """
            Finds needle: "needle" in list "haystack" and returns the index
            :param needle: what to find
            :param haystack: where to find it.
            :return: the index for needle in haystack
            """
            for index, sub_list in enumerate(haystack):
                if needle in sub_list:
                    return index
            # Needle not in haystack
            return None

        def find_next_host(needle, haystack1, haystack2, hoststack):
            """
            Searches for the the guest lists to see where a participant is supposed to go next.
            :param haystack1: list to search in
            :param haystack2: list to search in
            :param hoststack: list to find hosts in
            :param needle: Who to search for
            :return: the name of the host
            """
            found_index = find_index(needle, haystack1)
            if found_index is None:
                found_index = find_index(needle, haystack2)
            if found_index is not None:
                return hoststack[found_index]
            else:
                found_index = find_index(needle, hoststack)
            if found_index is not None:
                return None
            else:
                # should not be possible to be here.
                raise IndexError("Error: 1")  # Unexpected Error, participant not found in any lists.

        def add_paragraph_next_stop(document, participant, text, host=None):
            """
            Add a paragraph that sends the participants to their next stop
            :param document: the document object
            :param participant: participant as list
            :param host: the host, if the participant is not hosting the next meal self
            :param text: the text to print
            :return: the modified document object.
            """

            if host is not None:
                document.add_paragraph(
                    "{}\n{}\n{}\n{}".format(
                        participant[0],
                        text,
                        host[0],
                        host[1]),
                    style=document.styles["Body Text 2"]
                )
            else:
                document.add_paragraph(
                    "{}\n{}".format(
                        participant[0],
                        text),
                    style=document.styles["Body Text 2"]
                )
            return document

        # Letters that are sent to participants before the food rally starts
        # ==================================================================
        if part == "starter" or part == "main_course" or part == "desert":
            for i in range(0, len(content)):
                # Heading containing the name
                doc.add_paragraph(content[i][0],
                                  style=doc.styles["Body Text 3"])
                # The part of the meal to prepare and allergies if any.
                doc.add_paragraph(
                    "{}".format(self.lang["word_to_prepare"]),
                    style=doc.styles["Body Text 2"]
                )
                doc.add_paragraph(
                    "{}".format(self.lang["word_{}".format(part)]),
                    style=doc.styles["Body Text 2"]
                )

                allergies = self.get_allergies(part, i)
                doc.add_paragraph(
                    allergies,
                    style=doc.styles["Body Text 2"]
                )

                # tell people where to have starters (if not a starter host)
                if part == "main_course" or part == "desert":
                    # Find the host
                    if any(content[i][0] == participant[0] for participant in self.guest_s_1):
                        host_index = find_index(content[i][0], self.guest_s_1)
                    elif any(content[i][0] == participant[0] for participant in self.guest_s_2):
                        host_index = find_index(content[i][0], self.guest_s_2)
                    else:
                        # Should not be possible to be here
                        raise IndexError(self.lang["error_impossible_error"])
                    doc.add_paragraph(
                        self.lang["word_goto_starters"].format(self.host_s[host_index][0], self.host_s[host_index][1]),
                        style=doc.styles["Body Text 2"]
                    )
                doc.add_page_break()
            return doc

        # Letters that are sent to the participants to be distributed during the dinners
        # ==============================================================================

        elif part == "where_to_go":
            # Create letters for the starters
            for i in range(0, len(self.host_s)):

                # State which host should have this note.
                doc.add_paragraph(
                    self.lang["word_note_read_by"].format(self.host_s[i][0]),
                    style=doc.styles["Body Text"]
                )

                # header text
                doc.add_paragraph(
                    self.lang["word_leave_from_starters"],
                    style=doc.styles["Body Text 2"]
                )

                # send the host.
                # The host will not host another meal. so just pass it on.
                next_stop = find_next_host(self.host_s[i][0], self.guest_m_1, self.guest_m_2, self.host_m)
                doc = add_paragraph_next_stop(doc, self.host_s[i], self.lang["word_go_to"], next_stop)

                # Send guest 1
                next_stop = find_next_host(self.guest_s_1[i][0], self.guest_m_1, self.guest_m_2, self.host_m)
                # participant is hosting next meal.
                if next_stop is None:
                    doc = add_paragraph_next_stop(doc, self.guest_s_1[i], self.lang["word_go_to_is_host"])
                else:
                    doc = add_paragraph_next_stop(doc, self.guest_s_1[i], self.lang["word_go_to"], next_stop)

                # Send guest 2
                next_stop = find_next_host(self.guest_s_2[i][0], self.guest_m_1, self.guest_m_2, self.host_m)
                # participant is hosting next meal.
                if next_stop is None:
                    doc = add_paragraph_next_stop(doc, self.guest_s_2[i], self.lang["word_go_to_is_host"])
                else:
                    doc = add_paragraph_next_stop(doc, self.guest_s_2[i], self.lang["word_go_to"], next_stop)

                doc.add_page_break()

            # Create letters for the main course
            for i in range(0, len(self.host_s)):

                # State which host should have this note.
                doc.add_paragraph(
                    self.lang["word_note_read_by"].format(self.host_m[i][0]),
                    style=doc.styles["Body Text"]
                )

                # Header text
                doc.add_paragraph(
                    self.lang["word_leave_from_main_course"],
                    style=doc.styles["Body Text 2"]
                )

                # send the host.
                # The host will not host another meal. so just pass it on.
                next_stop = find_next_host(self.host_m[i][0], self.guest_d_1, self.guest_d_2, self.host_d)
                doc = add_paragraph_next_stop(doc, self.host_m[i], self.lang["word_go_to"], next_stop)

                # Send guest 1
                next_stop = find_next_host(self.guest_m_1[i][0], self.guest_d_1, self.guest_d_2, self.host_d)
                # participant is hosting next meal.
                if next_stop is None:
                    doc = add_paragraph_next_stop(doc, self.guest_m_1[i], self.lang["word_go_to_is_host"])
                else:
                    doc = add_paragraph_next_stop(doc, self.guest_m_1[i], self.lang["word_go_to"], next_stop)

                # Send guest 2
                next_stop = find_next_host(self.guest_m_2[i][0], self.guest_d_1, self.guest_d_2, self.host_d)
                # participant is hosting next meal.
                if next_stop is None:
                    doc = add_paragraph_next_stop(doc, self.guest_m_2[i], self.lang["word_go_to_is_host"])
                else:
                    doc = add_paragraph_next_stop(doc, self.guest_m_2[i], self.lang["word_go_to"], next_stop)

                # Don't add a page break after the final entry since this will be the last page.
                if i+1 != len(self.host_m):
                    doc.add_page_break()

            return doc

        else:
            # Should not be here, just return the document untouched
            return doc

    def get_allergies(self, part, i):
        """
        Get all allergies as a list.
        :param part: part of the meal
        :param i: index value
        :return: A string containing all allergies. or None if no allergies.
        """
        allergies = ""
        if part == "starter":
            participants = [self.host_s[i], self.guest_s_1[i], self.guest_s_2[i]]
        elif part == "main_course":
            participants = [self.host_m[i], self.guest_m_1[i], self.guest_m_2[i]]
        else:
            participants = [self.host_d[i], self.guest_d_1[i], self.guest_d_2[i]]

        for participant in participants:
            if participant[2] is not None:
                allergies += "{}, ".format(participant[2])
        if len(allergies) < 1:
            return "{}{}".format(self.lang["word_allergies"], self.lang["word_no_allergies"])
        return "{}{}".format(self.lang["word_allergies"], allergies[:-2])

    def generate_result(self):
        """
        Start the process of generating the results.
        """
        # Read the participant file contents.
        read_file_result = self.read_file_contents()
        if read_file_result == 0:
            if self.iv_new_year_same_lineup.get():
                self.log_output(self.lang["progress_prev_lineup_read"])
            else:
                self.log_output(self.lang["progress_lineup_read"])
        elif read_file_result == 1:
            self.log_output(self.lang["error_file_types"], "red")
            self.log_output(self.lang["error_not_excel"], "red")
            return
        else:
            self.log_output(self.lang["error_unexpected_read_file"], "red")
            return

        try:
            self.validate_number_of_participants()
        except ValueError as e:
            self.log_output("{}: {}".format(self.lang["error"], e), "red")
            return

        # Only create new groups of participants if previous line up not should be taken into account.
        if not self.iv_new_year_same_lineup.get():
            self.generate_random_index()
        self.sort_participants()

        # Generate a new lineup.
        self.create_lineup()

        self.log_output(self.lang["progress_done_saving"])

        # Save the result
        self.save_to_file()

        if self.iv_generate_letters.get():
            self.generate_docx_letters()

    def log_output(self, string, color="black"):
        """
        Method to print text to the output frame.
        :param string: The text to print
        :param color: Text color
        """
        self.t_output.configure(state=tkinter.NORMAL)
        if not string.endswith("\n"):
            string += "\n"
        self.t_output.insert(tkinter.END, string, color)
        self.t_output.yview_moveto(1)
        self.t_output.configure(state=tkinter.DISABLED)

    def get_lang(self):
        """
        Read a csv file and return the phrases matching the selected language
        self.lang is set to a dict with phrases in selected language
        """
        cur_lang = {}
        try:
            with open(self.csv_file, "r", encoding="utf8") as csv_file:
                reader = csv.DictReader(csv_file, delimiter=",")
                for row in reader:
                    cur_lang[row["phrase"]] = row[self.gui_language].replace("\\n", "\n")
        except KeyError:
            pass
        except FileNotFoundError:
            messagebox.showerror("File not found", "The language file: '{}' can not be found.\n"
                                 "Make sure it is available and try again.".format(self.csv_file))
            quit()
        self.lang = cur_lang

        # ============
        # END OF CLASS
        # ============

if __name__ == "__main__":
    root = tkinter.Tk()
    hmi = Hmi(root, language="swe")
    hmi.draw_main()
    root.mainloop()
